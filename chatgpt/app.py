import streamlit as st
import pandas as pd
from openai import OpenAI
from docx import Document
from io import BytesIO, StringIO

apiGroq = "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"
# Configuración de claves API (NO RECOMENDADO almacenar claves en código en producción)
client = OpenAI(base_url="https://api.groq.com/openai/v1", api_key=apiGroq)

#Funcion para generar articu
def generador_articulos(tema):
    try:
        response = client.chat.completions.create(
            model="llama-3.2-90b-vision-preview",
            messages=[
                {"role": "system", "content": "Eres un experto en generacion de articulos de SEO"},
                {"role": "user", "content": f"Escribe un articulo optimizado para SEO sobre {tema}"}
            ],
            max_tokens=1000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error al generar el articulo: {e}")
        return None

#Funcion para generar codigo
def generadorCodigo(descripcion):
    try:
        response = client.chat.completions.create(
            model="llama-3.2-90b-vision-preview",
            messages=[
                {"role": "system", "content": "Eres un programador de Python. Proporciona SOLO el codigo Python, sin explicaciones ni comentarios adicionales."},
                {"role": "user", "content": f"Escribe el codigo Python para: {descripcion} IMPORTANTE: Proporciona SOLO el codigo Python, sin explicaciones ni comentarios adicionales."}
            ],
            temperature=0.2
        )
        code = response.choices[0].message.content.strip()
        lineasCodigo = []
        inCodeBlock = False
        for linea in code.split("\n"):
            if linea.strip().startswith('```python'):
                inCodeBlock = True
                continue
            elif linea.strip().startswith('```'):
                inCodeBlock = False
                continue
            elif inCodeBlock or not linea.strip().startswith('```'):
                lineasCodigo.append(linea)

        return '\n'.join(lineasCodigo).strip()
    except Exception as e:
        st.error(f"Error al generar el codigo: {e}")
        return None
    
#Funcion para generar tablas de datos
def generadorTablas(descripcion):
    try:
        response = client.chat.completions.create(
            model="llama-3.2-90b-vision-preview",
            messages=[
                {"role": "system", "content": "Genera datos en formato CSV con encabezados. Los datos deben ser coherentes y bien estructurados. SOLO DEVUELVE EL TEXTO CSV NADA DE EXPLICACION"},   
                {"role": "user", "content": f"Genera una tabla de datos en formato CSV para: {descripcion}. Incluye encabezados y al menos 10 filas de datos "}
            ],
            temperature=0.5
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error al generar la tabla de datos: {e}")
        return None


#Funcion para generar un archivo de word
def crearWordDocumento(articulo):
    doc = Document()
    doc.add_heading('Articulo generado', 0)
    doc.add_paragraph(articulo)

    buff = BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff

#Funcion para generar un archivo de excel
def crearExcelDocumento(tabla):
    #Converi csv en dataframe
    lines = [line.strip() for line in tabla.split('\n') if line.strip()]
    if not lines:
        raise ValueError("El contenido de la tabla está vacío")

    #Separar encabezados y datos
    headers = lines[0].split(',')
    data_rows = [line.split(',') for line in lines[1:]]

    df = pd.DataFrame(data_rows, columns=headers)

    #Guardar excel con formato mejorado
    buff = BytesIO()
    #df.to_excel(buff, index=False, sheet_name='Datos', auto_filter=True)  # auto_filter para ajuste automático
    with pd.ExcelFile(buff, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False,sheet_name='Datos')
        #Ajustar el ancho de las columnas al contenido
        worksheet = writer.sheets['Datos']
        for idx, col in enumerate(df):  # loop through all columns
            max_length  = max(df[col].astype(str).map(len).max(), len(col))  # find the maximum length of the column name and its values
            worksheet.colum_dimensions[chr(65 + idx)].width = max_length + 2  # set the column width to the maximum length found
    buff.seek(0)
    return buff, df
        

#Configuracion de la pagina
def main():
    st.set_page_config(page_title="App",page_icon=":bar_chart:", 
                       layout="wide")
    st.title("Generador de contenidos con IA")
    seleccion = st.sidebar.selectbox("Seleccione una seccion:",
                             ["Generador de articulos", "Generacion de codigo","Generacion de tablas de datos"]
    )
    match seleccion:
        case "Generador de articulos":
            st.header("Generador de articulos")
            topic = st.text_input("Ingrese el tema del articulo")
            noEmpty = True if (topic.strip() and topic != "" ) else False
            if (st.button("Generar articulo") and noEmpty):
                with st.spinner("Generando articulo..."):
                    articulo = generador_articulos(topic)
                    if articulo:
                        st.success("Articulo generado con exito")
                        st.markdown("### Vista previa del articulo: ")
                        st.markdown(articulo)

                        #Boton de descarga
                        st.download_button(
                            "Descargar como Word",
                            data=crearWordDocumento(articulo),
                            file_name="articulo_{topic}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            else:
                st.warning("Por favor ingrese un tema")

        case "Generacion de codigo":
            st.header("Generacion de codigo")
            descripcion = st.text_input("Ingrese la descripcion del codigo")
            noEmpty = True if (descripcion.strip() and descripcion != "" ) else False
            if (st.button("Generar codigo") and noEmpty):
                with st.spinner("Generando codigo..."):
                    codigo = generadorCodigo(descripcion)
                    if codigo:
                        st.success("Codigo generado con exito")
                        st.markdown("### Vista previa del codigo: ")
                        st.code(codigo, language="python")
                        #Boton de descarga
                        st.download_button(
                            "Descargar como Python",
                            data = codigo,
                            file_name="Codigo.py",
                            mime = "text/x-python"
                        )
                    else:
                        st.error("Error al generar el codigo")
            else:
                st.warning("Por favor ingrese una descripcion")


        case "Generacion de tablas de datos":
            st.header("Generacion de tablas de datos")  
            descripcion = st.text_input("Ingrese la descripcion de la tabla de datos")
            if (st.button("Generar tabla")):
                with st.spinner("Generando tabla..."):
                    tabla = generadorTablas(descripcion)
                    if tabla:
                        st.success("Tabla generada con exito")
                        st.markdown("### Vista previa de la tabla: ")
                        tabladf = pd.read_csv(StringIO(tabla))
                        st.dataframe(tabladf)
                        #Boton de descarga
                        st.download_button(
                           "Descargar como CSV",
                           data = tabladf.to_csv(index=False),
                           file_name="Tabla.csv",
                           mime = "text/csv"
                        )
                    else:
                        st.error("Error al generar la tabla")
            else:
                st.warning("Por favor ingrese una descripcion")
            

if __name__ == "__main__":
    main()
